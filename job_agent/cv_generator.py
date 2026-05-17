import json
import re
from pathlib import Path

import anthropic

CV_BASE_NAMES = ['cv_base.pdf', 'cv_base.docx']


def load_cv_base(cv_dir="cv"):
    """Busca cv_base.pdf o cv_base.docx en cv_dir. Retorna (filepath, extension)."""
    for name in CV_BASE_NAMES:
        path = Path(cv_dir) / name
        if path.exists():
            return str(path), path.suffix.lower()
    raise FileNotFoundError(
        f"CV base no encontrado. Sube tu CV como '{cv_dir}/cv_base.pdf' o '{cv_dir}/cv_base.docx'."
    )


def extract_cv_text(filepath, ext):
    """Extrae texto plano del CV."""
    if ext == '.pdf':
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        return '\n'.join(page.extract_text() or '' for page in reader.pages)
    elif ext in ('.docx', '.doc'):
        from docx import Document
        doc = Document(filepath)
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        raise ValueError(f"Formato no soportado: {ext}")


# ---------------------------------------------------------------------------
# DOCX path: Claude identifies paragraphs by index and returns new text
# ---------------------------------------------------------------------------

def _adapt_docx_with_claude(paragraphs, oferta, api_key):
    """
    Recibe lista de (indice, texto) de todos los parrafos del DOCX.
    Claude identifica cuales adaptar y devuelve {indice, nuevo} pairs.
    """
    para_list = '\n'.join(f"[{i}] {text}" for i, text in paragraphs)

    prompt = f"""Eres un experto en CVs. Te doy una lista numerada con TODOS los parrafos de un CV y una oferta de empleo.

Tu tarea es identificar y adaptar SOLO estos elementos del CV:
1. El resumen/perfil profesional (parrafo largo tipo "about me", "sobre mi", o similar)
2. El titulo del puesto (ej: "Full-stack web developer", "Desarrollador full-stack")
3. La linea de skills/tecnologias (ej: "React JavaScript TypeScript ES6 HTML5...")
4. La descripcion del trabajo mas reciente (si existe un parrafo largo describiendo responsabilidades)

Reglas estrictas:
- NO cambies: educacion, datos personales (email, telefono, linkedin), idiomas, informacion adicional, fechas
- NO inventes experiencia ni habilidades que no esten en el CV original
- Adapta el lenguaje y el enfasis para que encaje con la oferta, destacando lo mas relevante
- Si un parrafo no necesita cambio, NO lo incluyas en la respuesta
- Devuelve el indice EXACTO de la lista

Responde UNICAMENTE con JSON valido (sin markdown):
{{
  "reemplazos": [
    {{"indice": 3, "nuevo": "nuevo texto adaptado"}},
    {{"indice": 7, "nuevo": "otro texto adaptado"}}
  ]
}}

PARRAFOS DEL CV:
{para_list}

OFERTA:
Titulo: {oferta.get('titulo', '')}
Empresa: {oferta.get('empresa', '')}
Descripcion: {oferta.get('descripcion', '')}"""

    client = anthropic.Anthropic(api_key=api_key)
    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )
    text = message.content[0].text.strip()
    text = re.sub(r'^```(?:json)?\s*', '', text)
    text = re.sub(r'\s*```$', '', text)
    result = json.loads(text)
    print(f"[cv_generator] Claude propone {len(result.get('reemplazos', []))} reemplazos")
    for r in result.get('reemplazos', []):
        print(f"  [indice={r['indice']}] {r['nuevo'][:80]}...")
    return result


def _replace_paragraph_text(para, new_text):
    """
    Reemplaza el texto de un parrafo conservando el estilo del primer Run.
    """
    first_run_style = {}
    if para.runs:
        r = para.runs[0]
        first_run_style = {
            'name': r.font.name,
            'size': r.font.size,
            'bold': r.bold,
            'italic': r.italic,
            'color': r.font.color.rgb if r.font.color and r.font.color.type else None,
        }

    for run in para.runs:
        run.text = ''

    run = para.runs[0] if para.runs else para.add_run()
    run.text = new_text

    if first_run_style.get('name'):
        run.font.name = first_run_style['name']
    if first_run_style.get('size'):
        run.font.size = first_run_style['size']
    if first_run_style.get('bold') is not None:
        run.bold = first_run_style['bold']
    if first_run_style.get('italic') is not None:
        run.italic = first_run_style['italic']
    if first_run_style.get('color'):
        run.font.color.rgb = first_run_style['color']


def _modify_docx(src_path, oferta, api_key, output_docx_path):
    """
    Abre el DOCX original, pide a Claude que identifique que parrafos adaptar
    (por indice), aplica los reemplazos conservando el formato, y guarda.
    """
    from docx import Document

    doc = Document(src_path)
    all_paragraphs = [(i, p.text) for i, p in enumerate(doc.paragraphs) if p.text.strip()]

    print(f"[cv_generator] DOCX tiene {len(doc.paragraphs)} parrafos, {len(all_paragraphs)} no vacios")

    result = _adapt_docx_with_claude(all_paragraphs, oferta, api_key)

    replaced = 0
    for reemplazo in result.get('reemplazos', []):
        idx = reemplazo.get('indice')
        nuevo = reemplazo.get('nuevo', '').strip()
        if idx is None or not nuevo:
            continue
        if 0 <= idx < len(doc.paragraphs):
            _replace_paragraph_text(doc.paragraphs[idx], nuevo)
            replaced += 1

    print(f"[cv_generator] {replaced} parrafos reemplazados en el DOCX")
    doc.save(output_docx_path)


# ---------------------------------------------------------------------------
# PDF path: extract text → Claude adapts → reportlab generates new PDF
# ---------------------------------------------------------------------------

def _adapt_pdf_with_claude(cv_text, oferta, api_key):
    """Para CV en PDF: extrae y adapta secciones, devuelve dict con nombre/contacto/resumen/etc."""
    client = anthropic.Anthropic(api_key=api_key)
    prompt = f"""Eres un experto en CVs. Dado el siguiente CV y una oferta de empleo, haz dos cosas:

1. Extrae del CV: nombre completo y datos de contacto (email, telefono, LinkedIn si existe).
2. Adapta SOLO estas secciones:
   - Resumen profesional (3-4 lineas enfocadas en la oferta)
   - Experiencia (selecciona y destaca los proyectos/logros mas relevantes; no inventes nada)
   - Skills (ordena y prioriza los que coincidan con la oferta)

NO cambies: educacion, certificaciones, datos personales.

Responde UNICAMENTE con JSON valido (sin markdown):
{{
  "nombre": "...",
  "contacto": "email | telefono | linkedin",
  "resumen": "...",
  "experiencia": ["bullet 1", "bullet 2", "..."],
  "skills": ["skill1", "skill2", "..."]
}}

CV BASE:
{cv_text}

OFERTA:
Titulo: {oferta.get('titulo', '')}
Empresa: {oferta.get('empresa', '')}
Descripcion: {oferta.get('descripcion', '')}"""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )
    text = message.content[0].text.strip()
    text = re.sub(r'^```(?:json)?\s*', '', text)
    text = re.sub(r'\s*```$', '', text)
    return json.loads(text)


def _render_pdf_from_scratch(adapted, output_path):
    """Genera un PDF profesional desde cero con reportlab (para CV en PDF)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, HRFlowable
    from reportlab.lib.enums import TA_CENTER

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    accent = colors.HexColor('#1a1a2e')
    style_name = ParagraphStyle('Name', fontSize=22, fontName='Helvetica-Bold',
                                alignment=TA_CENTER, spaceAfter=4)
    style_contact = ParagraphStyle('Contact', fontSize=10, fontName='Helvetica',
                                   alignment=TA_CENTER, textColor=colors.gray, spaceAfter=14)
    style_heading = ParagraphStyle('Heading', fontSize=11, fontName='Helvetica-Bold',
                                   textColor=accent, spaceBefore=14, spaceAfter=6)
    style_body = ParagraphStyle('Body', fontSize=10, fontName='Helvetica', leading=15, spaceAfter=4)
    style_bullet = ParagraphStyle('Bullet', fontSize=10, fontName='Helvetica',
                                  leading=14, leftIndent=14, spaceAfter=3)

    story = []
    story.append(Paragraph(adapted.get('nombre', ''), style_name))
    story.append(Paragraph(adapted.get('contacto', ''), style_contact))
    story.append(HRFlowable(width='100%', thickness=1.5, color=accent))
    story.append(Paragraph('RESUMEN PROFESIONAL', style_heading))
    story.append(Paragraph(adapted.get('resumen', ''), style_body))
    story.append(Paragraph('EXPERIENCIA', style_heading))
    for bullet in adapted.get('experiencia', []):
        story.append(Paragraph(f'• {bullet}', style_bullet))
    story.append(Paragraph('SKILLS', style_heading))
    story.append(Paragraph('  •  '.join(adapted.get('skills', [])), style_body))
    doc.build(story)


# ---------------------------------------------------------------------------
# LibreOffice conversion
# ---------------------------------------------------------------------------

def _docx_to_pdf(docx_path, pdf_path):
    """Convierte DOCX a PDF usando LibreOffice via subprocess."""
    import shutil
    import subprocess

    if not shutil.which('libreoffice'):
        raise RuntimeError(
            "LibreOffice no esta instalado. Instalalo con:\n"
            "  sudo apt-get install libreoffice"
        )

    out_dir = str(Path(pdf_path).parent)
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf', docx_path, '--outdir', out_dir],
        capture_output=True,
        text=True
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice fallo al convertir: {result.stderr.strip()}")

    expected = Path(out_dir) / (Path(docx_path).stem + '.pdf')
    if expected != Path(pdf_path):
        expected.rename(pdf_path)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_adapted_cv(oferta, user_config, cv_dir="cv"):
    """
    Genera un CV adaptado a la oferta y lo guarda como PDF.
    - DOCX input: Claude identifica parrafos por indice → modifica DOCX → LibreOffice → PDF
    - PDF input: extrae texto → Claude adapta secciones → reportlab genera PDF
    Retorna la ruta del PDF generado.
    """
    api_key = user_config.get('ANTHROPIC_API_KEY')
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY no configurada para este usuario.")

    filepath, ext = load_cv_base(cv_dir)
    print(f"[cv_generator] CV base: {filepath} ({ext})")

    Path(cv_dir).mkdir(exist_ok=True)
    output_pdf = str(Path(cv_dir) / f"generated_{oferta['id']}.pdf")

    if ext in ('.docx', '.doc'):
        tmp_docx = str(Path(cv_dir) / f"_tmp_{oferta['id']}.docx")
        try:
            _modify_docx(filepath, oferta, api_key, tmp_docx)
            _docx_to_pdf(tmp_docx, output_pdf)
        finally:
            Path(tmp_docx).unlink(missing_ok=True)
    else:
        cv_text = extract_cv_text(filepath, ext)
        adapted = _adapt_pdf_with_claude(cv_text, oferta, api_key)
        _render_pdf_from_scratch(adapted, output_pdf)

    return output_pdf
